"""
Generates AltairGO_Project_Manual.docx
Run: python generate_manual.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Default font ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Color palette ─────────────────────────────────────────────────────────────
INDIGO   = RGBColor(0x4F, 0x46, 0xE5)   # headings
DARK     = RGBColor(0x1E, 0x29, 0x3B)   # body
MUTED    = RGBColor(0x64, 0x74, 0x8B)   # captions
GREEN    = RGBColor(0x06, 0x6B, 0x39)   # code / success
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "4F46E5"   # table header fill (hex string for OxmlElement)
ALT_ROW   = "F1F5F9"   # alternating row fill


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def add_h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(26)
    run.font.color.rgb = INDIGO
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = INDIGO
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.color.rgb = color or DARK
    run.font.size      = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size      = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent   = Inches(0.4 * (level + 1))
    p.paragraph_format.space_after   = Pt(2)
    return p


def add_code(text):
    """Grey-background code block."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name       = 'Courier New'
    run.font.size       = Pt(9.5)
    run.font.color.rgb  = GREEN
    # Light grey shading via XML
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F1F5F9')
    shd.set(qn('w:val'),  'clear')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_table(headers, rows, col_widths=None):
    """Styled table with indigo header row and alternating body rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, HEADER_BG)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Body rows
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = ALT_ROW if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row_cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacing after table
    return table


def add_divider():
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '4F46E5')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
add_h1("AltairGO")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Project Manual")
run.font.size = Pt(18)
run.font.color.rgb = MUTED

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Complete guide for running the AltairGO travel intelligence platform")
run2.font.size = Pt(11)
run2.italic    = True
run2.font.color.rgb = MUTED

doc.add_paragraph()
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_h2("1. What Is AltairGO?")
add_body(
    "AltairGO is an AI-powered travel planning platform. You tell it where you want to go, "
    "your budget, and how many days you have — and it builds a complete day-by-day itinerary "
    "with real cost estimates, hotel suggestions, activities, restaurants, and more. "
    "It can also help you book everything in one click and gives you a day-of briefing "
    "so you always know what to expect."
)

add_h3("The Two Parts of AltairGO")
add_table(
    ["Part", "Folder on Your Computer", "What It Does"],
    [
        ["AltairGO Engine\n(Backend)", "D:\\Projects\\AltairGO-Engine-main",
         "The brain — handles all logic, AI, database, calculations, and bookings"],
        ["AltairGO Platform\n(Frontend)", "D:\\Projects\\AltairGO-Platform",
         "The face — the website that users open in their browser"],
    ],
    col_widths=[1.8, 2.8, 2.8]
)

add_body(
    "Think of it like a restaurant: the Engine is the kitchen doing all the work, "
    "the Platform is the dining room the customer sees. Both must be running at the same time.",
    italic=True, color=MUTED
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ONE-TIME INSTALLS
# ══════════════════════════════════════════════════════════════════════════════
add_h2("2. Things to Install (One-Time Only)")
add_body("Download and install these programs before anything else:")
add_table(
    ["Program", "Where to Download", "Why You Need It"],
    [
        ["Python 3.10+",  "python.org",     "Runs the Engine (backend)"],
        ["Node.js 18+",   "nodejs.org",     "Runs the Platform (frontend)"],
        ["Docker Desktop","docker.com",     "Runs Redis — a fast memory store the app needs"],
        ["Git",           "Already installed", "Version control (you already have this)"],
    ],
    col_widths=[1.5, 2.2, 3.5]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. API KEYS
# ══════════════════════════════════════════════════════════════════════════════
add_h2("3. Accounts and API Keys You Need")
add_body(
    "An API key is like a password that lets AltairGO use a service. "
    "All of these are FREE."
)
add_table(
    ["Service", "Cost", "What It Does", "How to Get It"],
    [
        ["Supabase",     "Free",     "Your database — stores all users, trips, destinations", "Go to supabase.com → create a free project → copy the connection string"],
        ["Google Gemini","Free tier","AI that writes itinerary descriptions in natural language","Go to aistudio.google.com → click Get API Key"],
        ["Pexels",       "Free",     "Photos for destinations and attractions","Go to pexels.com/api → sign up free → copy your API key"],
        ["Redis",        "Free",     "Fast memory cache (runs inside Docker)","No account needed — Docker handles it automatically"],
    ],
    col_widths=[1.4, 0.8, 2.6, 2.5]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. FIRST-TIME SETUP
# ══════════════════════════════════════════════════════════════════════════════
add_h2("4. First-Time Setup (Do This Once)")
add_body("Follow these steps in order. You only need to do this one time.")

add_h3("Step 1 — Start Docker Desktop")
add_body(
    "Open Docker Desktop from your Start menu or taskbar. "
    "Wait until the whale icon appears in your system tray and stops animating — "
    "it is then ready. You do not need to do anything else inside Docker Desktop."
)

add_h3("Step 2 — Set Up the Engine (Backend)")
add_body("Open a terminal (Command Prompt or PowerShell) and navigate to the Engine folder:")
add_code("cd D:\\Projects\\AltairGO-Engine-main")
add_body("Create a private Python environment:")
add_code("python -m venv .venv")
add_body("Activate it (Windows):")
add_code(".venv\\Scripts\\activate")
add_body("You will see (.venv) appear at the start of the line — this means it is active. Now install all packages:")
add_code("pip install -r backend/requirements.txt")
add_body("This will download packages and takes 1-3 minutes. Wait for it to finish.")

add_h3("Step 3 — Create Your Settings File")
add_body("Still in the Engine folder, copy the example settings file:")
add_code("copy .env.example .env")
add_body(
    "Now open the .env file in Notepad (or any text editor) and fill in your values:"
)
add_code(
    "DATABASE_URL=<paste your Supabase connection string here>\n"
    "REDIS_URL=redis://localhost:6379/0\n"
    "JWT_SECRET_KEY=<type any random 40+ character string — mash the keyboard>\n"
    "ADMIN_ACCESS_KEY=<a password you will use for the admin panel>\n"
    "GEMINI_API_KEY=<from aistudio.google.com>\n"
    "PEXELS_API_KEY=<from pexels.com/api>\n"
    "DEV_EAGER=true"
)

add_body(
    "Where to find the Supabase connection string: "
    "Log into supabase.com → open your project → Settings → Database → "
    "Connection String → select URI mode → copy the full string.",
    italic=True, color=MUTED
)

add_h3("Step 4 — Set Up the Platform (Frontend)")
add_body("Open a SECOND terminal and navigate to the Platform folder:")
add_code("cd D:\\Projects\\AltairGO-Platform")
add_body("Install all packages:")
add_code("npm install")
add_body("This takes 1-2 minutes. Done — first-time setup is complete.")
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 5. SEED THE DATABASE
# ══════════════════════════════════════════════════════════════════════════════
add_h2("5. Populate Your Database (First Time Only)")
add_body(
    "Your database starts empty. Run these commands once to fill it with "
    "destinations, attractions, prices, and settings. "
    "Make sure the Engine is running (Step 6) before you do this."
)
add_body("In the Engine terminal (with .venv active), run each command and wait for it to finish:")

steps = [
    ("1", "Ingest attractions from OpenStreetMap (free map data)", "python -m backend.scripts.ingest_osm_data"),
    ("2", "Enrich with descriptions and images",                  "python -m backend.scripts.enrich_attractions"),
    ("3", "Build geospatial index",                               "python -m backend.scripts.h3_indexer"),
    ("4", "Calculate popularity scores",                          "python -m backend.scripts.score_attractions"),
    ("5", "Add hotel and flight pricing",                         "python -m backend.scripts.sync_prices"),
    ("6", "Set engine settings to defaults",                      "python -m backend.scripts.init_settings"),
]
for num, desc, cmd in steps:
    add_body(f"Step {num} — {desc}:")
    add_code(cmd)
add_body("Each command will print progress messages. Wait for each to finish before running the next one.")
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 6. STARTING THE APP (EVERY DAY)
# ══════════════════════════════════════════════════════════════════════════════
add_h2("6. Starting the App Every Day")
add_body(
    "Every time you want to use AltairGO, you need three things running at the same time. "
    "Use three separate terminal windows."
)

add_h3("Terminal 1 — Start Redis")
add_body("Open a terminal and run:")
add_code("docker run -d -p 6379:6379 redis:7")
add_body(
    "You only need to run this once per day (or configure Docker Desktop to keep Redis always on). "
    "Once it prints a long ID and exits, Redis is running in the background.",
    italic=True, color=MUTED
)

add_h3("Terminal 2 — Start the Engine (Backend)")
add_body("Navigate to the Engine folder and activate your Python environment:")
add_code("cd D:\\Projects\\AltairGO-Engine-main")
add_code(".venv\\Scripts\\activate")
add_body("Start the backend server:")
add_code("python -m flask --app backend.app:create_app run --port 5000 --reload")
add_body("You will see:")
add_code("* Running on http://127.0.0.1:5000")
add_body("Leave this terminal open. Do not close it while using the app.")

add_h3("Terminal 3 — Start the Platform (Frontend)")
add_body("Open a new terminal, navigate to the Platform folder:")
add_code("cd D:\\Projects\\AltairGO-Platform")
add_body("Start the website:")
add_code("npm run dev")
add_body("You will see:")
add_code("Local: http://localhost:5173")

add_h3("Open Your Browser")
add_body(
    "Go to http://localhost:5173 — your AltairGO app is live!",
    bold=True
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CHEAT SHEET
# ══════════════════════════════════════════════════════════════════════════════
add_h2("7. Daily Cheat Sheet — Quick Reference")
add_table(
    ["Step", "Terminal", "Command"],
    [
        ["1. Start Docker Desktop",       "—",          "Open Docker Desktop from Start menu"],
        ["2. Start Redis",                "Terminal 1", "docker run -d -p 6379:6379 redis:7"],
        ["3. Activate Python env",        "Terminal 2", "cd D:\\Projects\\AltairGO-Engine-main\n.venv\\Scripts\\activate"],
        ["4. Start Engine",               "Terminal 2", "python -m flask --app backend.app:create_app run --port 5000 --reload"],
        ["5. Start Platform",             "Terminal 3", "cd D:\\Projects\\AltairGO-Platform\nnpm run dev"],
        ["6. Open browser",               "Browser",    "http://localhost:5173"],
    ],
    col_widths=[2.0, 1.4, 4.0]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 8. PAGES GUIDE
# ══════════════════════════════════════════════════════════════════════════════
add_h2("8. What Each Page Does")
add_table(
    ["Page Name", "Web Address", "What It Does"],
    [
        ["Home",            "localhost:5173/",                 "Landing page — intro to AltairGO"],
        ["Discover",        "localhost:5173/discover",         "Browse and filter destinations"],
        ["Destination",     "localhost:5173/destination/123",  "Detailed info about a specific place"],
        ["Trip Planner",    "localhost:5173/planner",          "Fill a form to generate a full AI itinerary"],
        ["My Trips",        "localhost:5173/trips",            "See all your saved trips (login required)"],
        ["Trip Viewer",     "localhost:5173/trip/123",         "Full itinerary — bookings, expenses, tools"],
        ["Daily Briefing",  "localhost:5173/trip/123/briefing/1","Day-of guide: what to wear, carry, and expect"],
        ["Profile",         "localhost:5173/profile",          "Update your travel preferences"],
        ["Blogs",           "localhost:5173/blogs",            "Travel blog articles"],
        ["Admin Panel",     "localhost:5173/admin",            "Backend control panel (admin key required)"],
    ],
    col_widths=[1.8, 2.8, 2.7]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 9. LOGGING IN
# ══════════════════════════════════════════════════════════════════════════════
add_h2("9. Logging In")

add_h3("As a Traveler (Regular User)")
add_bullet("Go to localhost:5173/register")
add_bullet("Fill in your name, email, and a password (minimum 12 characters)")
add_bullet("You are now logged in — your trips will save to your account")
add_bullet("Next time: go to localhost:5173/login")

add_h3("As Admin")
add_bullet("Go to localhost:5173/admin/login")
add_bullet("Enter the ADMIN_ACCESS_KEY you set in your .env file")
add_bullet("You now have access to the admin dashboard — user management, destination control, engine settings")
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 10. HOW TO GENERATE A TRIP
# ══════════════════════════════════════════════════════════════════════════════
add_h2("10. How to Generate a Trip (Step-by-Step)")
steps_trip = [
    "Go to localhost:5173/planner",
    "Select your destination country and city",
    "Choose your destinations (e.g. Jaipur, Udaipur)",
    "Set your total budget in Indian Rupees (e.g. 15000)",
    "Set number of days (e.g. 3) and number of travelers",
    "Pick your travel style: Budget / Balanced / Luxury",
    "Select traveler type: Solo / Couple / Family / Friends",
    "Choose travel month (e.g. December)",
    "Optionally set dietary restrictions, accessibility needs, or children count",
    "Click Generate Itinerary",
    "Wait 5-15 seconds — the AI builds your complete day-by-day plan",
    "Review your itinerary, save it, and start booking!",
]
for i, s in enumerate(steps_trip, 1):
    add_bullet(f"{i}. {s}")
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 11. RUNNING TESTS
# ══════════════════════════════════════════════════════════════════════════════
add_h2("11. Checking That Everything Is Working (Tests)")
add_body(
    "The Engine comes with 188 automated tests that verify everything is working correctly. "
    "Run them any time you want to confirm the backend is healthy."
)
add_body("In the Engine terminal (with .venv active):")
add_code("python -m pytest backend/tests/ -q")
add_body("You should see:")
add_code("188 passed, 1 skipped")
add_body(
    "If you see failures, check that your .env file is filled in correctly "
    "and that the database connection is working.",
    italic=True, color=MUTED
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 12. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
add_h2("12. Troubleshooting — Common Problems")
add_table(
    ["Problem You See", "What to Do"],
    [
        ["Cannot connect to database",
         "Open your .env file and check DATABASE_URL. Copy a fresh connection string from Supabase → Settings → Database → URI."],
        ["Redis connection refused",
         "Make sure Docker Desktop is open and running. Then run: docker run -d -p 6379:6379 redis:7"],
        ["No module named ...",
         "You forgot to activate your Python environment. Run: .venv\\Scripts\\activate (in the Engine folder)"],
        ["Frontend shows a blank white page",
         "Make sure the Engine (Terminal 2) is running on port 5000. Check Terminal 2 for error messages."],
        ["Invalid token / you get logged out immediately",
         "Your JWT_SECRET_KEY may have changed in .env. Log out, clear browser storage, and log back in."],
        ["Port 5000 already in use",
         "Another program is using port 5000. Start Flask on a different port: add --port 5001 to the command."],
        ["npm run dev fails",
         "Navigate to the Platform folder first: cd D:\\Projects\\AltairGO-Platform — then run npm install again."],
        ["Itinerary generation hangs forever",
         "Check that DEV_EAGER=true is in your .env file. This setting lets the Engine work without a separate worker process."],
    ],
    col_widths=[3.0, 4.3]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 13. GITHUB
# ══════════════════════════════════════════════════════════════════════════════
add_h2("13. GitHub Repositories")
add_table(
    ["Repository", "URL", "Contents"],
    [
        ["AltairGO Engine",   "github.com/yash-dev007/AltairGO-Engine",   "Backend API, pipeline, tests, scripts"],
        ["AltairGO Platform", "github.com/yash-dev007/AltairGO-Platform", "React frontend website"],
    ],
    col_widths=[2.0, 3.2, 2.1]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AltairGO Engine — Built by yash-dev007")
run.font.size = Pt(9)
run.font.color.rgb = MUTED
run.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "AltairGO_Project_Manual.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
